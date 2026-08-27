"""Evrak üretimi ve teslim takibi testleri."""

from __future__ import annotations

from datetime import date, timedelta
from pathlib import Path

import pytest

from cekirdek.modeller import PlanParametreleri
from evrak import uretici
from evrak.belge import Belge, tr_tarih
from veri import hizmet
from veri.hizmet import HizmetHatasi
from veri.veritabani import Veritabani
from testler.test_hizmet import AYARLAR, _kur


@pytest.fixture()
def hazir(tmp_path: Path):
    """Kaydedilmiş planı olan bir veritabanı."""
    vt = Veritabani(tmp_path / "sorumluluk.db")
    vt.gocleri_uygula()
    _kur(vt, tmp_path)
    plan_id = hizmet.plan_kaydet(vt, hizmet.plan_hazirla(
        vt, PlanParametreleri(pencere_kodu="P1")))
    return vt, plan_id, tmp_path


def _metin(yol: Path) -> str:
    """Belgenin görünen metnini üstbilgi ve altbilgi dâhil toplar."""
    from docx import Document
    belge = Document(yol)
    parcalar = []
    for bolum in belge.sections:
        for parca in (bolum.header, bolum.footer):
            parcalar.extend(p.text for p in parca.paragraphs)
    parcalar.extend(p.text for p in belge.paragraphs)
    for tablo in belge.tables:
        for satir in tablo.rows:
            parcalar.extend(h.text for h in satir.cells)
    return "\n".join(parcalar)


# =============================================================== belge katmanı

def test_tarih_turkce_bicimde_yazilir() -> None:
    assert tr_tarih(date(2026, 9, 14)) == "14.09.2026"
    assert tr_tarih("2026-09-14") == "14.09.2026"
    assert tr_tarih(None) == ""


def test_ayni_icerik_ayni_ozeti_verir(tmp_path: Path) -> None:
    """Sürüm numarası içerik özetine bağlıdır; .docx bayt karşılaştırması
    güvenilir bir ölçüt değildir."""
    def kur():
        b = Belge("Deneme Lisesi", "Deneme", "alt")
        b.paragraf("Bir paragraf.")
        b.tablo(["A", "B"], [("1", "2")])
        return b
    assert kur().kaydet(tmp_path / "a.docx") == kur().kaydet(tmp_path / "b.docx")
    farkli = kur()
    farkli.paragraf("fazladan")
    assert farkli.kaydet(tmp_path / "c.docx") != kur().kaydet(tmp_path / "d.docx")


def test_bos_tablo_kayit_yok_yazar(tmp_path: Path) -> None:
    b = Belge("Deneme", "Deneme")
    b.tablo(["A", "B"], [])
    b.kaydet(tmp_path / "bos.docx")
    assert "Kayıt yok" in _metin(tmp_path / "bos.docx")


# ============================================================== evrak üretimi

def test_tum_evraklar_uretilir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    uretilenler = uretici.evrak_uret(vt, plan_id, tmp_path / "evrak")
    assert len(uretilenler) == len(uretici.EVRAKLAR)
    for yol, ozet in uretilenler:
        assert yol.exists() and yol.stat().st_size > 0
        assert len(ozet) == 64


def test_secili_evrak_uretilir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    uretilenler = uretici.evrak_uret(vt, plan_id, tmp_path / "evrak",
                                     ["01_sinav_programi", "03_gorevlendirme_cizelgesi"])
    assert [y.name for y, _ in uretilenler] == ["01_sinav_programi.docx",
                                                "03_gorevlendirme_cizelgesi.docx"]


def test_bilinmeyen_evrak_turu_reddedilir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    with pytest.raises(HizmetHatasi, match="Bilinmeyen evrak türü"):
        uretici.evrak_uret(vt, plan_id, tmp_path / "evrak", ["99_olmayan"])


def test_sinav_programi_oturumlari_ve_kurumu_icerir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "program.docx"
    uretici.sinav_programi(vt, plan_id, yol)
    metin = _metin(yol)
    assert AYARLAR["okul_adi"] in metin
    assert "MATEMATİK" in metin
    assert "TASLAK" in metin           # plan henüz kesinleşmedi


def test_kesinlesen_planin_programinda_onay_no_yazar(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    hizmet.plan_kesinlestir(vt, plan_id, "2026/144")
    yol = tmp_path / "program.docx"
    uretici.sinav_programi(vt, plan_id, yol)
    metin = _metin(yol)
    assert "2026/144" in metin
    assert "TASLAK" not in metin


def test_gorevli_nushasi_gorevlileri_icerir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    ogrenci_nushasi = tmp_path / "ogrenci.docx"
    gorevli_nushasi = tmp_path / "gorevli.docx"
    uretici.sinav_programi(vt, plan_id, ogrenci_nushasi, gorevli_nushasi=False)
    uretici.sinav_programi(vt, plan_id, gorevli_nushasi, gorevli_nushasi=True)
    assert "Komisyon üyesi" in _metin(gorevli_nushasi)
    assert "Komisyon üyesi" not in _metin(ogrenci_nushasi)


def test_gorevlendirme_oluru_dayanagi_tasir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "olur.docx"
    uretici.gorevlendirme_cizelgesi(vt, plan_id, yol)
    metin = _metin(yol)
    assert "md.58/2-a" in metin
    assert "Karar md.12/2-b" in metin
    assert AYARLAR["mudur_adi"] in metin


def test_gorev_sayac_raporu_ucret_hesaplamaz(hazir) -> None:
    """Kapsam gereği tutar hesaplanmaz; yalnız sayaç tutulur."""
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "sayac.docx"
    uretici.gorev_sayac_raporu(vt, plan_id, yol)
    metin = _metin(yol)
    assert "Komisyon" in metin and "Gözcülük" in metin
    assert "md.12/2-a" in metin
    assert "TL" not in metin and "brüt" not in metin.lower()


# ============================================================ sürüm takibi

def test_ayni_icerik_yeni_surum_acmaz(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    uretici.evrak_uret(vt, plan_id, tmp_path / "evrak")
    ilk = len(hizmet.evrak_gecmisi(vt, str(plan_id)))
    uretici.evrak_uret(vt, plan_id, tmp_path / "evrak")
    assert len(hizmet.evrak_gecmisi(vt, str(plan_id))) == ilk


def test_icerik_degisince_yeni_surum_acilir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    uretici.evrak_uret(vt, plan_id, tmp_path / "evrak", ["01_sinav_programi"])
    with vt.baglan() as b:
        b.execute("UPDATE oturum SET saat='15:30' WHERE plan_id=?", (plan_id,))
    uretici.evrak_uret(vt, plan_id, tmp_path / "evrak", ["01_sinav_programi"])
    surumler = [s for t, s, _, _ in hizmet.evrak_gecmisi(vt, str(plan_id))
                if t == "01_sinav_programi"]
    assert surumler == [1, 2]


# ============================================================ teslim takibi

def test_cizelge_her_oturum_icin_beklenen_evraki_listeler(hazir) -> None:
    vt, plan_id, _ = hazir
    cizelge = hizmet.teslim_cizelgesi(vt, plan_id)
    oturum_sayisi = len(hizmet.plan_oturumlari(vt, plan_id))
    assert len(cizelge) == oturum_sayisi * len(hizmet.BEKLENEN_EVRAK)
    assert all(not s.teslim_edildi_mi for s in cizelge)


def test_teslim_kaydedilir_ve_cizelgeye_yansir(hazir) -> None:
    vt, plan_id, _ = hazir
    oturum = hizmet.plan_oturumlari(vt, plan_id)[0]
    personel = [p for p in hizmet.personelleri_getir(vt) if p.gorev_alabilir_mi]
    hizmet.teslim_kaydet(vt, oturum["id"], "sinav_kagitlari",
                         personel[0].kimlik, personel[1].kimlik, adet=12,
                         aciklama="Zarf içinde")
    satir = next(s for s in hizmet.teslim_cizelgesi(vt, plan_id)
                 if s.oturum_id == oturum["id"] and s.evrak_turu == "sinav_kagitlari")
    assert satir.teslim_edildi_mi
    assert satir.adet == 12
    assert satir.teslim_eden == personel[0].ad
    assert satir.durum() == "teslim alındı"


def test_ts03_teslim_eden_ve_alan_ayni_kisi_olamaz(hazir) -> None:
    vt, plan_id, _ = hazir
    oturum = hizmet.plan_oturumlari(vt, plan_id)[0]
    kisi = hizmet.personelleri_getir(vt)[0]
    with pytest.raises(HizmetHatasi, match="aynı kişi olamaz"):
        hizmet.teslim_kaydet(vt, oturum["id"], "sinav_kagitlari",
                             kisi.kimlik, kisi.kimlik)


def test_ts02_suresi_gecen_evrak_gecikmis_sayilir(hazir) -> None:
    vt, plan_id, _ = hazir
    cizelge = hizmet.teslim_cizelgesi(vt, plan_id)
    satir = cizelge[0]
    # Teslim süresi sınav tarihini izleyen ilk iş günüdür.
    assert not satir.gecikti_mi(satir.tarih)
    assert satir.gecikti_mi(satir.son_gun() + timedelta(days=1))
    assert satir.durum(satir.son_gun() + timedelta(days=1)) == "gecikti"


def test_teslim_edilen_evrak_gecikmis_sayilmaz(hazir) -> None:
    vt, plan_id, _ = hazir
    oturum = hizmet.plan_oturumlari(vt, plan_id)[0]
    personel = [p for p in hizmet.personelleri_getir(vt) if p.gorev_alabilir_mi]
    hizmet.teslim_kaydet(vt, oturum["id"], "komisyon_tutanagi",
                         personel[0].kimlik, personel[1].kimlik)
    satir = next(s for s in hizmet.teslim_cizelgesi(vt, plan_id)
                 if s.oturum_id == oturum["id"] and s.evrak_turu == "komisyon_tutanagi")
    assert not satir.gecikti_mi(date(2099, 1, 1))


def test_teslim_geri_alinabilir(hazir) -> None:
    vt, plan_id, _ = hazir
    oturum = hizmet.plan_oturumlari(vt, plan_id)[0]
    personel = [p for p in hizmet.personelleri_getir(vt) if p.gorev_alabilir_mi]
    hizmet.teslim_kaydet(vt, oturum["id"], "sinav_kagitlari",
                         personel[0].kimlik, personel[1].kimlik)
    hizmet.teslim_geri_al(vt, oturum["id"], "sinav_kagitlari")
    satir = next(s for s in hizmet.teslim_cizelgesi(vt, plan_id)
                 if s.oturum_id == oturum["id"] and s.evrak_turu == "sinav_kagitlari")
    assert not satir.teslim_edildi_mi


def test_teslim_ozeti_sayilari_toplar(hazir) -> None:
    vt, plan_id, _ = hazir
    oturum = hizmet.plan_oturumlari(vt, plan_id)[0]
    personel = [p for p in hizmet.personelleri_getir(vt) if p.gorev_alabilir_mi]
    hizmet.teslim_kaydet(vt, oturum["id"], "sinav_kagitlari",
                         personel[0].kimlik, personel[1].kimlik)
    ozet = hizmet.teslim_ozeti(vt, plan_id, bugun=date(2026, 9, 1))
    assert ozet["teslim"] == 1
    assert ozet["toplam"] == ozet["teslim"] + ozet["bekliyor"] + ozet["gecikti"]


def test_gecersiz_evrak_turu_reddedilir(hazir) -> None:
    vt, plan_id, _ = hazir
    oturum = hizmet.plan_oturumlari(vt, plan_id)[0]
    personel = hizmet.personelleri_getir(vt)
    with pytest.raises(HizmetHatasi, match="Geçersiz evrak türü"):
        hizmet.teslim_kaydet(vt, oturum["id"], "olmayan_evrak",
                             personel[0].kimlik, personel[1].kimlik)


def test_teslim_tutanagi_gecikmeyi_gosterir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "teslim.docx"
    uretici.evrak_teslim_tutanagi(vt, plan_id, yol, bugun=date(2099, 1, 1))
    metin = _metin(yol)
    assert "gecikti" in metin
    assert "süresinde teslim edilmemiştir" in metin


def test_baslik_turkce_buyutulur(tmp_path: Path) -> None:
    """Yerleşik upper() 'i' harfini bozar: LISTESI yazardı."""
    b = Belge("Deneme", "Sınav Yoklama ve Salon Listesi")
    b.kaydet(tmp_path / "baslik.docx")
    metin = _metin(tmp_path / "baslik.docx")
    assert "SINAV YOKLAMA VE SALON LİSTESİ" in metin


# ============================================ komisyon bazlı görevlendirme

def test_gorevlendirme_cizelgesi_komisyon_bazlidir(hazir) -> None:
    """Her satır bir sınav komisyonudur; kişi başına satır açılmaz."""
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "cizelge.docx"
    uretici.gorevlendirme_cizelgesi(vt, plan_id, yol)
    from docx import Document
    tablolar = Document(yol).tables
    basliklar = [h.text for h in tablolar[0].rows[0].cells]
    assert basliklar == ["Sınav Yapılacak Ders veya Dersler", "Öğrenci\nSayısı",
                         "Sınav\nTarihi", "Sınav\nSaati", "Sınav Salonu",
                         "Komisyon Üyeleri", "Gözcü / Gözcüler"]
    # Başlık dışında satır sayısı oturum sayısına eşit olmalı.
    assert len(tablolar[0].rows) - 1 == len(hizmet.plan_oturumlari(vt, plan_id))


def test_gorevlendirme_cizelgesinde_teblig_tebellug_bolumu_var(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "cizelge.docx"
    uretici.gorevlendirme_cizelgesi(vt, plan_id, yol)
    from docx import Document
    belge = Document(yol)
    metin = _metin(yol)
    assert "TEBLİĞ - TEBELLÜĞ BELGESİ" in metin
    assert "Tebliğ Tarihi" in metin and "İmza" in metin
    # Tebliğ tablosunda görevli her personel bir kez yer alır.
    gorevliler = hizmet.gorevli_listesi(vt, plan_id)
    assert len(belge.tables[1].rows) - 1 == len(gorevliler)


def test_birlestirilmis_dersler_cizelgede_belirtilir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "cizelge.docx"
    uretici.gorevlendirme_cizelgesi(vt, plan_id, yol)
    metin = _metin(yol)
    birlesik = [o for o in hizmet.plan_oturumlari(vt, plan_id) if "/" in o["duzey"]]
    if birlesik:
        assert "birleştirilmiştir" in metin


def test_komisyon_uyeleri_ayri_satirlarda_yazilir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "cizelge.docx"
    uretici.gorevlendirme_cizelgesi(vt, plan_id, yol)
    from docx import Document
    hucre = Document(yol).tables[0].rows[1].cells[5]
    assert len(hucre.paragraphs) == 2          # iki komisyon üyesi, iki paragraf


# ==================================================== KVKK ilan çizelgeleri

def test_ilan_belgelerinde_hicbir_kisi_adi_gecmez(hazir) -> None:
    """İlan çıktıları herkese açık bir sayfada yayımlanır: ne öğrenci, ne
    görevli öğretmen, ne de imzalayan müdürün adı yazılır. Belgeyi çıkaran
    makam kişi adı olmadan gösterilir."""
    vt, plan_id, tmp_path = hazir
    takvim = tmp_path / "ilan_takvim.docx"
    cizelge = tmp_path / "ilan_ogrenci.docx"
    uretici.ilan_sinav_takvimi(vt, plan_id, takvim)
    uretici.ilan_ogrenci_cizelgesi(vt, plan_id, cizelge)
    for yol in (takvim, cizelge):
        metin = _metin(yol)
        for kisi in hizmet.personelleri_getir(vt, yalniz_aktif=False):
            assert kisi.ad not in metin, f"{yol.name} içinde {kisi.ad} geçiyor"
        assert AYARLAR["mudur_adi"] not in metin
        assert "Okul Müdürlüğü" in metin
        assert "6698" in metin


def test_ilan_takviminde_ogrenci_verisi_yok(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "ilan_takvim.docx"
    uretici.ilan_sinav_takvimi(vt, plan_id, yol)
    metin = _metin(yol)
    for ogrenci in hizmet.sorumluluk_kayitlari(vt):
        assert ogrenci.ad_soyad not in metin
        assert ogrenci.okul_no not in metin
    assert "MATEMATİK" in metin


def test_ilan_ogrenci_cizelgesinde_acik_ad_yazmaz(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yol = tmp_path / "ilan_ogrenci.docx"
    uretici.ilan_ogrenci_cizelgesi(vt, plan_id, yol)
    metin = _metin(yol)
    for ogrenci in hizmet.sorumluluk_kayitlari(vt):
        assert ogrenci.ad_soyad not in metin
    assert "U****** Ö****** B**" in metin      # maskeli ad
    assert "101" in metin                      # okul numarası
    assert "6698" in metin


def test_ilan_cizelgesinde_gosterim_secenegi_uygulanir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    yalniz_no = tmp_path / "no.docx"
    yalniz_ad = tmp_path / "ad.docx"
    uretici.ilan_ogrenci_cizelgesi(vt, plan_id, yalniz_no, "yalniz_no")
    uretici.ilan_ogrenci_cizelgesi(vt, plan_id, yalniz_ad, "yalniz_maskeli_ad")
    assert "U******" not in _metin(yalniz_no)
    assert "U******" in _metin(yalniz_ad)


def test_gecersiz_gosterim_bicimi_reddedilir(hazir) -> None:
    vt, plan_id, _ = hazir
    with pytest.raises(HizmetHatasi, match="Geçersiz öğrenci gösterim"):
        hizmet.ilan_ogrenci_cizelgesi(vt, plan_id, "acik_ad")


def test_ilan_ogrenci_cizelgesi_ogrenci_basina_gruplanir(hazir) -> None:
    vt, plan_id, _ = hazir
    cizelge = hizmet.ilan_ogrenci_cizelgesi(vt, plan_id)
    assert len(cizelge) == 3                   # üç öğrenci
    ilk = next(o for o in cizelge if o["okul_no"] == "101")
    assert len(ilk["sinavlar"]) == 2           # matematik + fizik


def test_taslak_plan_ilan_belgelerinde_uyari_tasir(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    for uretici_fn, ad in ((uretici.ilan_sinav_takvimi, "takvim"),
                           (uretici.ilan_ogrenci_cizelgesi, "cizelge")):
        yol = tmp_path / f"{ad}.docx"
        uretici_fn(vt, plan_id, yol)
        assert "TASLAK" in _metin(yol)


def test_kesinlesen_planin_ilaninda_uyari_yok(hazir) -> None:
    vt, plan_id, tmp_path = hazir
    hizmet.plan_kesinlestir(vt, plan_id, "2026/144")
    yol = tmp_path / "takvim.docx"
    uretici.ilan_sinav_takvimi(vt, plan_id, yol)
    assert "TASLAK" not in _metin(yol)
