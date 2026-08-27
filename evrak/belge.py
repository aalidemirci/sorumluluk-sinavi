"""Resmî evrak için .docx yapı taşları.

Belgeler şablon dosyasından değil doğrudan koddan üretilir: depoda ikili
dosya durmaz, şablon ile kod birbirinden kopmaz ve sayfa düzeni tek yerde
tanımlanır.

Sayfa düzeni Türk resmî yazışma alışkanlığına göredir: A4 dikey, 2,5 cm
kenar boşluğu, üstbilgide kurum adı, altbilgide sayfa numarası.
"""

from __future__ import annotations

import hashlib
from datetime import date
from pathlib import Path

from cekirdek.metin import buyult


LACIVERT = "17365D"
BASLIK_ZEMINI = "E8EEF5"
SATIR_ZEMINI = "F7F9FB"
CIZGI = "AAB6C4"
SOLUK = "5C6773"
KIRMIZI = "9B1C1C"


def tr_tarih(deger: date | str | None) -> str:
    """Tarihi gg.aa.yyyy biçiminde yazar."""
    if deger is None or deger == "":
        return ""
    if isinstance(deger, str):
        try:
            deger = date.fromisoformat(deger[:10])
        except ValueError:
            return deger
    return deger.strftime("%d.%m.%Y")


class Belge:
    """Tek bir resmî evrakı kuran yardımcı."""

    def __init__(self, kurum_adi: str, baslik: str, alt_baslik: str = "",
                 yatay: bool = False):
        from docx import Document
        from docx.enum.section import WD_ORIENT
        from docx.shared import Mm, Pt, RGBColor

        self.belge = Document()
        bolum = self.belge.sections[0]
        if yatay:
            bolum.orientation = WD_ORIENT.LANDSCAPE
            bolum.page_width, bolum.page_height = Mm(297), Mm(210)
        else:
            bolum.orientation = WD_ORIENT.PORTRAIT
            bolum.page_width, bolum.page_height = Mm(210), Mm(297)
        bolum.left_margin = bolum.right_margin = Mm(25)
        bolum.top_margin = bolum.bottom_margin = Mm(20)
        bolum.header_distance = bolum.footer_distance = Mm(12)

        normal = self.belge.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(10.5)
        normal.font.color.rgb = RGBColor.from_string("20252B")
        normal.paragraph_format.space_before = Pt(0)
        normal.paragraph_format.space_after = Pt(5)
        normal.paragraph_format.line_spacing = 1.1

        self._ustbilgi(kurum_adi)
        self._altbilgi()
        self._baslik(baslik, alt_baslik)

    # ------------------------------------------------------------ düzen

    def _yazi(self, run, boyut=10.5, kalin=False, renk="20252B", italik=False):
        from docx.oxml.ns import qn
        from docx.shared import Pt, RGBColor
        run.font.name = "Calibri"
        rpr = run._element.get_or_add_rPr()
        rpr.get_or_add_rFonts().set(qn("w:ascii"), "Calibri")
        rpr.rFonts.set(qn("w:hAnsi"), "Calibri")
        run.font.size = Pt(boyut)
        run.bold = kalin
        run.italic = italik
        run.font.color.rgb = RGBColor.from_string(renk)
        return run

    def _kenarlik(self, paragraf, yer="bottom", renk=CIZGI, kalinlik="6"):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        ppr = paragraf._p.get_or_add_pPr()
        bdr = ppr.find(qn("w:pBdr"))
        if bdr is None:
            bdr = OxmlElement("w:pBdr")
            ppr.append(bdr)
        el = bdr.find(qn("w:" + yer))
        if el is None:
            el = OxmlElement("w:" + yer)
            bdr.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), kalinlik)
        el.set(qn("w:space"), "4")
        el.set(qn("w:color"), renk)

    @staticmethod
    def _golgele(nesne, renk):
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        pr = (nesne._p.get_or_add_pPr() if hasattr(nesne, "_p")
              else nesne._tc.get_or_add_tcPr())
        shd = pr.find(qn("w:shd"))
        if shd is None:
            shd = OxmlElement("w:shd")
            pr.append(shd)
        shd.set(qn("w:fill"), renk)

    def _ustbilgi(self, kurum_adi: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        paragraf = self.belge.sections[0].header.paragraphs[0]
        paragraf.text = kurum_adi
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraf.paragraph_format.space_after = Pt(2)
        self._kenarlik(paragraf, "bottom", CIZGI, "4")
        for run in paragraf.runs:
            self._yazi(run, 8.5, True, SOLUK)

    def _altbilgi(self) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt
        paragraf = self.belge.sections[0].footer.paragraphs[0]
        paragraf.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        paragraf.paragraph_format.space_before = Pt(3)
        self._kenarlik(paragraf, "top", CIZGI, "4")
        self._yazi(paragraf.add_run("Sayfa "), 8, False, SOLUK)
        alan = OxmlElement("w:fldSimple")
        alan.set(qn("w:instr"), "PAGE")
        paragraf._p.append(alan)

    def _baslik(self, baslik: str, alt_baslik: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        paragraf = self.belge.add_paragraph()
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraf.paragraph_format.space_after = Pt(3)
        # Yerleşik upper() Türkçede 'i' harfini bozar: LISTESI / LİSTESİ.
        self._yazi(paragraf.add_run(buyult(baslik)), 14, True, LACIVERT)
        self._kenarlik(paragraf, "bottom", LACIVERT, "12")
        if alt_baslik:
            alt = self.belge.add_paragraph()
            alt.alignment = WD_ALIGN_PARAGRAPH.CENTER
            alt.paragraph_format.space_after = Pt(10)
            self._yazi(alt.add_run(alt_baslik), 9.5, False, SOLUK)

    def yeni_bolum_basligi(self, baslik: str, alt_baslik: str = '') -> None:
        """Sayfa sonrasında başlığı yeniden yazar; çok sayfalı evraklar için."""
        self._baslik(baslik, alt_baslik)

    # ------------------------------------------------------------ içerik

    def paragraf(self, metin: str, kalin: bool = False, boyut: float = 10.5,
                 renk: str = "20252B", ortala: bool = False, bosluk: int = 5):
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        paragraf = self.belge.add_paragraph()
        paragraf.paragraph_format.space_after = Pt(bosluk)
        if ortala:
            paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self._yazi(paragraf.add_run(metin), boyut, kalin, renk)
        return paragraf

    def bilgi_satirlari(self, ciftler: list[tuple[str, str]]) -> None:
        """Etiket/değer çiftlerini iki sütunlu, çerçevesiz bir bloğa yazar."""
        from docx.shared import Pt
        for etiket, deger in ciftler:
            paragraf = self.belge.add_paragraph()
            paragraf.paragraph_format.space_after = Pt(2)
            self._yazi(paragraf.add_run(f"{etiket}: "), 10, True)
            self._yazi(paragraf.add_run(str(deger)), 10)

    def tablo(self, basliklar: list[str], satirlar: list[tuple],
              genislikler: list[int] | None = None, bos_metin: str = "Kayıt yok"):
        """Başlık satırı gölgeli, kenarlıklı bir tablo ekler.

        `genislikler` yirmide bir punto (dxa) cinsindendir; verilmezse
        sütunlar eşit bölünür.
        """
        from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        from docx.shared import Pt

        sutun_sayisi = len(basliklar)
        tablo = self.belge.add_table(rows=1, cols=sutun_sayisi)
        tablo.alignment = WD_TABLE_ALIGNMENT.LEFT
        tablo.autofit = False

        bolum = self.belge.sections[0]
        kullanilabilir = int((bolum.page_width - bolum.left_margin - bolum.right_margin) / 635)
        if not genislikler or len(genislikler) != sutun_sayisi:
            taban = kullanilabilir // sutun_sayisi
            genislikler = [taban] * (sutun_sayisi - 1) + [kullanilabilir - taban * (sutun_sayisi - 1)]
        else:
            toplam = sum(genislikler)
            genislikler = [int(g * kullanilabilir / toplam) for g in genislikler]

        pr = tablo._tbl.tblPr
        tblw = OxmlElement("w:tblW")
        tblw.set(qn("w:w"), str(kullanilabilir))
        tblw.set(qn("w:type"), "dxa")
        pr.append(tblw)
        bordur = OxmlElement("w:tblBorders")
        for yon in ("top", "left", "bottom", "right", "insideH", "insideV"):
            el = OxmlElement("w:" + yon)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), "4")
            el.set(qn("w:color"), CIZGI)
            bordur.append(el)
        pr.append(bordur)
        for sutun, genislik in zip(tablo._tbl.tblGrid.gridCol_lst, genislikler):
            sutun.set(qn("w:w"), str(genislik))

        for hucre, baslik in zip(tablo.rows[0].cells, basliklar):
            hucre.text = baslik
        for veri in (satirlar or [(bos_metin,) + ("",) * (sutun_sayisi - 1)]):
            hucreler = tablo.add_row().cells
            for hucre, deger in zip(hucreler, list(veri) + [""] * sutun_sayisi):
                hucre.text = "" if deger is None else str(deger)

        for satir_no, satir in enumerate(tablo.rows):
            trpr = satir._tr.get_or_add_trPr()
            trpr.append(OxmlElement("w:cantSplit"))
            if satir_no == 0:
                tekrar = OxmlElement("w:tblHeader")
                tekrar.set(qn("w:val"), "true")
                trpr.append(tekrar)
            for sutun_no, (hucre, genislik) in enumerate(zip(satir.cells, genislikler)):
                hucre.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
                self._golgele(hucre, LACIVERT if satir_no == 0
                              else (SATIR_ZEMINI if satir_no % 2 == 0 else "FFFFFF"))
                tcw = hucre._tc.get_or_add_tcPr().get_or_add_tcW()
                tcw.set(qn("w:w"), str(genislik))
                tcw.set(qn("w:type"), "dxa")
                for paragraf in hucre.paragraphs:
                    paragraf.paragraph_format.space_before = Pt(1)
                    paragraf.paragraph_format.space_after = Pt(1)
                    paragraf.paragraph_format.line_spacing = 1.0
                    if satir_no == 0 or sutun_no == 0:
                        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraf.runs:
                        self._yazi(run, 8.5 if sutun_sayisi >= 7 else 9.5,
                                   satir_no == 0, "FFFFFF" if satir_no == 0 else "20252B")
        self.belge.add_paragraph()
        return tablo

    def dayanak_notu(self, metin: str) -> None:
        from docx.shared import Pt
        paragraf = self.belge.add_paragraph()
        paragraf.paragraph_format.space_before = Pt(6)
        self._yazi(paragraf.add_run(metin), 8.5, False, SOLUK, italik=True)

    def imza_blogu(self, sutunlar: list[tuple[str, str]], olur_adi: str = "") -> None:
        """Sekme duraklarıyla imza bloğu kurar; tablo kullanılmaz."""
        from docx.enum.text import WD_TAB_ALIGNMENT
        from docx.shared import Mm, Pt

        girisler = list(sutunlar) + ([("OLUR", olur_adi)] if olur_adi else [])
        if not girisler:
            return
        bolum = self.belge.sections[0]
        genislik = bolum.page_width - bolum.left_margin - bolum.right_margin
        adim = genislik / len(girisler)

        roller = self.belge.add_paragraph()
        roller.paragraph_format.space_before = Pt(20)
        roller.paragraph_format.keep_together = True
        adlar = self.belge.add_paragraph()
        adlar.paragraph_format.space_before = Pt(26)
        for paragraf in (roller, adlar):
            for sira in range(1, len(girisler)):
                paragraf.paragraph_format.tab_stops.add_tab_stop(
                    Mm(int((adim * sira + adim / 2) / 36000)), WD_TAB_ALIGNMENT.CENTER)
        self._yazi(roller.add_run("\t".join(rol for rol, _ in girisler)), 9.5, True, LACIVERT)
        self._yazi(adlar.add_run("\t".join(ad or "…………………………" for _, ad in girisler)), 9)

    def uyari(self, metin: str) -> None:
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
        paragraf = self.belge.add_paragraph()
        paragraf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        paragraf.paragraph_format.space_after = Pt(8)
        self._golgele(paragraf, "FDE8E7")
        self._yazi(paragraf.add_run(metin), 9.5, True, KIRMIZI)

    def sayfa_sonu(self) -> None:
        from docx.enum.text import WD_BREAK
        self.belge.add_paragraph().add_run().add_break(WD_BREAK.PAGE)

    # ------------------------------------------------------------- kayıt

    def icerik_ozeti(self) -> str:
        """Belgenin görünen metninden SHA-256 özeti üretir.

        Dosya baytları kullanılamaz: .docx bir zip arşividir ve içindeki zaman
        damgaları her üretimde değiştiği için aynı içerik farklı bayt özeti
        verir. Sürüm numarası bu yüzden metne bağlanır — içerik değişmediyse
        yeni sürüm açılmaz.
        """
        parcalar: list[str] = []
        for bolum in self.belge.sections:
            for parca in (bolum.header, bolum.footer):
                parcalar.extend(p.text for p in parca.paragraphs)
        parcalar.extend(p.text for p in self.belge.paragraphs)
        for tablo in self.belge.tables:
            for satir in tablo.rows:
                parcalar.extend(h.text for h in satir.cells)
        metin = "\n".join(" ".join(p.split()) for p in parcalar)
        return hashlib.sha256(metin.encode("utf-8")).hexdigest()

    def kaydet(self, hedef: Path) -> str:
        """Belgeyi yazar ve içerik özetini döndürür."""
        hedef = Path(hedef)
        hedef.parent.mkdir(parents=True, exist_ok=True)
        ozet = self.icerik_ozeti()
        self.belge.save(hedef)
        return ozet
